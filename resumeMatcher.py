import requests
from bs4 import BeautifulSoup
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_minimal_line_spacing(paragraph):
    paragraph.paragraph_format.line_spacing = 1  # Single line spacing
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)

def printDocument(allSkills, certifications, workExperience, projects, schoolCourses):
    doc = Document()
    name = "First Last"
    number = "347-333-3335"
    email = "name2@rpi.edu"
    location = "New York, NY"
    linkedin = "https://www.linkedin.com/in/firstlast/"
    github = "https://github.com/firstlast"
    schoolName = "Rensselaer Polytechnic Institute"
    schoolLocation = "Troy, NY"
    schoolGraduation = "May 2025"
    schoolMajor = "Computer Science"
    majorConcentration = "Finance"

    # Personal Information
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    set_minimal_line_spacing(p)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{number}, {email}, {location}")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    set_minimal_line_spacing(p)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{linkedin}, {github}")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    set_minimal_line_spacing(p)

    # Skills and Certifications Section
    p = doc.add_paragraph()
    run = p.add_run("Skills: ")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run = p.add_run(', '.join(allSkills))
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    set_minimal_line_spacing(p)

    p = doc.add_paragraph()
    run = p.add_run("Certifications: ")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run = p.add_run(', '.join(certifications))
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    set_minimal_line_spacing(p)

    # Work Experience Section
    for experience in workExperience:
        company_info = f"{experience[0]}, {experience[1]}"
        work_date = experience[2]
        spaces_needed = 100 - len(company_info)  # Adjust the number of spaces as needed
        line_text = company_info + ' ' * spaces_needed + work_date

        p = doc.add_paragraph(line_text)
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
        set_minimal_line_spacing(p)

        p = doc.add_paragraph()
        run = p.add_run(experience[3])
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        set_minimal_line_spacing(p)

        for bullet in experience[4:]:
            p = doc.add_paragraph(bullet)
            run = p.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            set_minimal_line_spacing(p)

    # Projects Section
    for project in projects:
        project_info = f"{project[0]}, {project[1]}"
        project_date = project[2]
        spaces_needed = 100 - len(project_info)  # Adjust the number of spaces as needed
        line_text = project_info + ' ' * spaces_needed + project_date

        p = doc.add_paragraph(line_text)
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
        set_minimal_line_spacing(p)

        for bullet in project[3:]:
            p = doc.add_paragraph(bullet)
            run = p.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            set_minimal_line_spacing(p)

    # Education Section
    p = doc.add_paragraph()
    run = p.add_run(f"EDUCATION\n{schoolName}, {schoolLocation}, {schoolGraduation}\n{schoolMajor}\nConcentration in {majorConcentration}\nRelevant Courses: {', '.join(schoolCourses)}")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    set_minimal_line_spacing(p)

    # Save the document
    doc.save("test.docx")
def scrape_all_words_from_page(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    words = re.findall(r'[a-zA-Z0-9-_.]+', soup.get_text().lower())
    return set(words)

def main():
    url = "https://careers.doordash.com/jobs/5511447"
    page_words_set = scrape_all_words_from_page(url)
    targetSet = { "git", "javascript", "react", "angular", "data structures", "embedded systems", "go", "rust", "C#", "C++", "C", "golang", "java","algorithms", "database", "SQL", "object oriented programming", "collaboration", "documentation", "backend", "distributed systems", "docker", "azure", "gcp", "linux", "rust", "jira", "XML", "html", "css", "typescript", "php", "tailwind", "bootstrap", "jQuery", "ajax", "vue", "node", "oracle", "postgresql", "mongodb", "mysql", "express", "springboot", "django", "flask", "graphql", "ux", "rest apis", "google cloud platform", "swift", "react native", "ruby", "python", "kotlin", "dart", "sql", "nosql", "postgres", "redis", "elasticsearch", "rabbitmq"}

    commons = page_words_set.intersection(targetSet)
    print("Major words:", commons)

    resumeSkills = {"JavaScript", "Node", "Express", "Firebase", "MongoDB", "Git",
    "React", "HTML", "C++", "Python", "TailwindCSS", "jQuery", "Back End",
    "MySQL", "Front End", "Azure", "Cloud Computing", "Google Cloud Platform",
    "AWS", "Rest APIs", "JavaFX", "Docker", "Rust", "Go","sql","java"}
    skills_strings = {skill.lower() for skill in resumeSkills}
    certifications = ["Oracle Cloud Infrastructure 2023 Foundations Associate"]

    percentMatch = len(commons.intersection(skills_strings)) / len(commons) if commons else 0
    print("Percent match: {:.2%}".format(percentMatch))

    desired_match_pct = 0.85
    if percentMatch >= desired_match_pct:
        print("The resume passes the desired ATS threshold.")
    else:
        extraWords = commons - skills_strings
        print("Extra Words:", extraWords)
        print("Length of Extra Words:", len(extraWords), "Length of Commons:", len(commons))

        extra_skills_needed = int(desired_match_pct * len(targetSet)) - len(commons.intersection(skills_strings))
        missing_skills = commons - skills_strings
        skills_to_add = list(missing_skills)[:extra_skills_needed]
        finalSkills = skills_to_add + list(resumeSkills)
        schoolCourses = ["Computer Science 1",
    "Data Structures",
    "Discrete Structures",
    "Computer Organization",
    "Intro to Algorithms",
    "Principles of Software",
    "Intro to IT and Web Science",
    "Intro to UX",
    "Web Systems Development",
    "Web Science Systems Development",
    "Managing IT Resources",
    "Database Systems"]
    workExperience = [["Facebook","Remote","May 2023-Current","Software Engineer Intern","• Led fullstack software development projects, utilizing a diverse tech stack including React, HTML/CSS, Node.js, Express.js, MongoDB, Firestore and more, with 5-15% of errors in the software product during the testing phase.","• Utilized a Cloud-based data storage platform to support the aggregation of data from multiple sources and the generation of reports, increasing the efficiency of the reporting process by 90% and enabling the team to focus on more critical tasks."],["Google","New York, NY","June 2022-August 2022", "Code Next Coach Intern","•   Helped the program director to facilitate a startup incubator program with a cohort of 17 in the Spring of 2018, with a 100% graduation rate and a 90% participant retention rate for students.","• Co-facilitated a full-stack development program utilizing Python and HTML/CSS/JS for a group of 10 students over 10 weeks, resulting in a final project showcasing students' understanding of the fundamentals of web development."],["IBM Accelerate Software Track","Remote","June 2023-July 2023","Student","•   React.JS, JavaScript, Back-End Development, Databases, Web Security and Cloud Native Development"]]
    projects = [["DNK Events, Gaddy Web Consulting","June 2023-July 2023","• Developed a profit-oriented, user-friendly website for DNK Events, leveraging technologies such as Firebase, HTML, CSS, and JavaScript, which resulted in a 25% increase in clientele from 1,017 to 1,271.","• Built a dynamic web page using Firebase Cloud Functions, Firebase Firestore, Bootstrap, Node.js, JavaScript, and Webflow, which boosted client engagement by 20%.","•   Designed and implemented an internal password-protected page for efficient customer management and consultation booking."],["Haute-Femme","Gaddy Web Consulting","July 2023-August 2023","• Converted HTML/CSS template to React.js and built e-commerce site using MERN stack (MongoDB, Express, React, Node.js), resulting in the client being able to manage their website remotely and efficiently in order to increase store revenue by implementing an e-commerce feature including cart and checkout using Stripe payment processing while concurrently handling 300+ monthly users."],["Housing 2.0","Open Source (Rensselear Open Source Club)","• Engineered a full-scale UI/UX revamp for the Housing 2.0 portal, resulting in a 35% increase in user engagement and a 50% faster navigation experience for 7000+ students.","•  Spearheaded scalability optimizations, enabling the platform to handle 3 times the original user load during critical operation times, ensuring seamless access to housing resources.","•   Innovated roommate and dorm matching algorithms, enhancing match accuracy and elevating student satisfaction by 25%, validated through extensive unit testing with 90% code coverage."],["PrograMeet","Independent","•  Designed and created multiple iterations of social media app that allows academically and career driven 1000+ high school students to connect (React Native, UI/UX utilizing Adobe XD, and metrics, surveys, and A/B testing). ","• Created and led the implementation of a database design in Firebase.","•    Received $1,000 in app funding from RPI business competition."]]
    printDocument(finalSkills, certifications, workExperience, projects, schoolCourses)

if __name__ == "__main__":
    main()