# import re
# import requests
# from bs4 import BeautifulSoup
# from docx import Document

# # Function to check if the text contains any of the targeted keywords
# def contains(targeted_keywords, text):
#     lower_text = text.lower()
#     return any(re.search(r'\b{}\b'.format(re.escape(keyword.lower())), lower_text) for keyword in targeted_keywords)

# # Function to scrape a website for skills
# def scrape_website(url, targeted_keywords):
#     response = requests.get(url)
#     soup = BeautifulSoup(response.content, 'html.parser')
#     found_skills = set()

#     # Search for the skills in the webpage text
#     text = soup.get_text(separator=' ')
#     for keyword in targeted_keywords:
#         if contains(targeted_keywords, text):
#             found_skills.add(keyword)

#     return list(found_skills)

# # Function to update the docx file with the skills
# def update_docx(file_path, skills):
#     doc = Document(file_path)
#     for para in doc.paragraphs:
#         if "SKILLS" in para.text:
#             # Assuming skills are listed in a line after the "SKILLS" heading
#             skills_line = para.text.split(":")[1] if ":" in para.text else ""
#             existing_skills = skills_line.split(", ")
#             # Add new skills that are not already in the document
#             new_skills = [skill for skill in skills if skill not in existing_skills]
#             para.text += ", " + ", ".join(new_skills)
#             break

#     doc.save(file_path)

# # Main function to run the script
# def main():
#     # Prompt user for the URL
#     url = input("What is the url of the site you'd like to query? ")
#     targeted_keywords = ["javascript", "nodejs", "expressjs", "firebase", "mongodb", "git",
#                          "react.js", "html", "c++", "python", "tailwindcss", "jquery", "back end",
#                          "mysql", "front end", "azure", "cloud computing", "google cloud platform",
#                          "aws", "rest apis", "javafx", "docker", "rust", "golang", "problem solving",
#                          "analytical", "reasoning", "communication", "collaborative"] 

#     # Scrape the website for skills
#     found_skills = scrape_website(url, targeted_keywords)

#     # File path to the .docx resume file
#     file_path = 'v1.docx'  # Replace with the path to your .docx file

#     # Update the .docx file with found skills
#     update_docx(file_path, found_skills)

#     print("The .docx file has been updated with the following skills:")
#     print(found_skills)

# if __name__ == "__main__":
#     main()

from docx import Document

def read_and_update_docx(file_path):
    try:
        # Attempt to open the .docx file
        doc = Document(file_path)
        # ... (the rest of your code for updating the docx file)
    except Exception as e:
        print(f"An error occurred: {e}")
        # If the document is corrupted or not a valid .docx, handle it here

# Replace 'path/to/your/document.docx' with the actual path to your .docx file
read_and_update_docx('copy.docx')
